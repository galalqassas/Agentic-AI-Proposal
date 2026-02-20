"""Chainlit UI for the Proposal Agent System.

Visualises the multi-agent workflow with correctly ordered, collapsible
steps nested under parent messages.  Only the final accepted proposal
appears as a standalone chat message.
"""

import io
import os
import re
from datetime import datetime, timezone

import chainlit as cl
from langchain_core.messages import HumanMessage

from data_layer import JsonDataLayer
from graph.graph import build_graph, QUALITY_THRESHOLD, MAX_ITERATIONS

# ── Initialisation ───────────────────────────────────────────────────
APP_GRAPH = build_graph()

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Nodes whose events we handle in the stream
MAJOR_NODES = frozenset({"planner", "researcher", "writer", "evaluator", "output", "ask_user"})

# Human-readable step names
_STEP_LABELS = {
    "planner": "Planner",
    "researcher": "Researcher",
    "output": "Output",
}


def _writer_label(n: int) -> str:
    return f"Writer (attempt {n})"


def _evaluator_label(n: int) -> str:
    return f"Evaluator (attempt {n})"


# ── Helpers ──────────────────────────────────────────────────────────
def _utc_now() -> str:
    """ISO-8601 UTC timestamp string."""
    return datetime.now(timezone.utc).isoformat()


def _score_emoji(score: float) -> str:
    if score >= 9.0:
        return "🟢"
    if score >= 7.0:
        return "🟡"
    return "🔴"


def _build_scorecard(
    dimension_scores: dict,
    overall_score: float,
    revision_count: int,
) -> str:
    rows = "\n".join(
        f"| {_score_emoji(v)} | {dim} | **{v}** |"
        for dim, v in dimension_scores.items()
    )
    emoji = _score_emoji(overall_score)
    return (
        f"## {emoji} Evaluation Scorecard\n\n"
        f"| | Dimension | Score |\n"
        f"|---|-----------|-------|\n"
        f"{rows}\n\n"
        f"**Overall: {overall_score}/10** · Revision {revision_count}/{MAX_ITERATIONS}"
    )


async def _end_step(step: cl.Step) -> None:
    """Mark a step as finished and push the update to the UI."""
    step.end = _utc_now()
    await step.update()


async def _make_step(name: str, parent_id: str) -> cl.Step:
    """Create, register, and send a step nested under *parent_id*."""
    step = cl.Step(name=name, type="run")
    step.parent_id = parent_id
    await step.send()
    return step


# ── Starters ─────────────────────────────────────────────────────────
@cl.set_starters
async def set_starters():
    return [
        cl.Starter(label="📋 Grant Proposal",       message="Act as Sarah Al-Fayed, Country Director at MSF. Write a grant proposal titled 'Operation Clean Water: Yemen 2026' to the Bill & Melinda Gates Foundation for a $750K emergency cholera intervention. Proposed start date: March 1st, 2026. Contact: s.alfayed@msf.org / +967-1-234567. Provide 5,000 cholera kits, deploy 12 medical staff, and initiate a 6-month budget (80% medical, 20% logistics)."),
        cl.Starter(label="💼 Business Plan",         message="Act as Alex Chen, Founder of LedgerLoop (Series A Fintech). Write a business plan pitching \"Stripe for Corporate Bonds\" to Sequoia Capital. Use a tech stack based on Rust/Solana; detail a 12-month roadmap (Q1-Q4); and justify a $15M ask broken down into 60% R&D, 30% Ops, and 10% Marketing."),
        cl.Starter(label="⚙️ Technical Proposal",    message="Act as TechFlow Solutions (AWS Partner). Write a technical proposal for First Midwest Bank to migrate from on-prem mainframes to AWS Cloud. Scale: 1500 VMs and 100TB database. Propose a phased \"6 Rs\" framework; guarantee SOC2 Type II compliance; and detail a zero-downtime cutover strategy."),
        cl.Starter(label="💰 Sales Proposal",        message="Act as a Salesforce Enterprise AE. Write a closing proposal for Mayo Clinic to adopt Health Cloud. Target a 15% reduction in patient wait times; include a 12-month contract; and present tiered pricing for 1,000 seats including a \"Co-Innovation Lab\" partnership."),
        cl.Starter(label="📅 Project Proposal",      message="Act as ThoughtWorks (Agile Dev Shop). Write a project proposal to build the MVP for \"NeoBank\". Include a 20-week total project duration with sprint-based delivery for a mobile app and admin dashboard; define the \"Definition of Done\" for the MVP; and use an embedded client product owner model."),
        cl.Starter(label="🔬 Research Proposal",     message="Act as Pfizer Oncology R&D. Write a Phase 3 Clinical Trial Protocol for Drug-X (Lung Cancer). Study 1,200 patients over 36 months; define primary endpoints for Overall Survival vs. Progression-Free Survival; and detail the global site selection and DSMB governance strategy."),
        cl.Starter(label="🤝 Partnership Proposal",  message="Act as Spotify Business Development. Write a partnership proposal to Uber. Include cross-platform authentication and weekly co-marketing syncs; propose a 10% revenue share on bookings; and focus on the technical API integration allowing riders to control the car stereo via Spotify."),
        cl.Starter(label="📝 General Proposal",      message="Act as Jessica Pearson, VP of People. Write an internal proposal to the Board of Directors for a 12-week \"4-Day Work Week\" pilot. KPIs: 20% increase in productivity, 15% reduction in turnover; trial the Monday-Thursday schedule in Engineering; and address accountability measures."),
    ]


# ── Auth & data layer ────────────────────────────────────────────────

@cl.header_auth_callback
async def header_auth_callback(headers: dict) -> cl.User:
    """Auto-authenticate every visitor — no login page."""
    return cl.User(identifier="default", metadata={"role": "admin"})


@cl.data_layer
def get_data_layer():
    return JsonDataLayer()


# ── Chat lifecycle ───────────────────────────────────────────────────

@cl.on_chat_start
async def start():
    cl.user_session.set("graph", APP_GRAPH)
    cl.user_session.set("config", {"configurable": {"thread_id": cl.context.session.id}})
    cl.user_session.set("processed_ids", set())
    cl.user_session.set("awaiting_planner_answers", False)
    cl.user_session.set("is_processing", False)


@cl.on_chat_resume
async def on_chat_resume(thread: dict):
    """Restore session state when a user reopens a past conversation."""
    cl.user_session.set("graph", APP_GRAPH)
    cl.user_session.set("config", {"configurable": {"thread_id": thread["id"]}})
    cl.user_session.set("processed_ids", set())
    cl.user_session.set("awaiting_planner_answers", False)
    cl.user_session.set("is_processing", False)


@cl.on_message
async def main(message: cl.Message):
    """Stream graph execution with steps nested under parent messages."""
    if cl.user_session.get("is_processing"):
        return

    cl.user_session.set("is_processing", True)
    try:
        graph = cl.user_session.get("graph")
        config = cl.user_session.get("config")

        state = graph.get_state(config)

        if state.next:
            # Check if we're resuming from planner questions (ask_user node).
            # We use a session flag because after ask_user interrupts,
            # state.next is ("researcher",) — not ("ask_user",).
            if cl.user_session.get("awaiting_planner_answers"):
                cl.user_session.set("awaiting_planner_answers", False)
                # User answered planner questions — merge into task
                existing_task = state.values.get("task", "")
                updated_task = f"{existing_task}\n\nAdditional info from user: {message.content}"
                await graph.aupdate_state(
                    config,
                    {"task": updated_task, "questions_for_user": []},
                    as_node="ask_user",
                )
            else:
                # Handle the "Proceed" button case vs actual typed feedback
                feedback = message.content
                intent = cl.user_session.get("intent")
                as_node = "researcher"  # Default: update researcher’s feedback but move to writer

                if feedback.startswith("✅ Proceeding"):
                    feedback = ""
                elif intent == "reresearch":
                    # Special case: Loop back to researcher node by updating as 'planner'
                    as_node = "planner"
                    cl.user_session.set("intent", None)
                elif intent == "edit":
                    # Regular case: feedback is for the writer, proceed normally
                    cl.user_session.set("intent", None)

                await graph.aupdate_state(
                    config, {"user_feedback": feedback}, as_node=as_node
                )
            stream = graph.astream_events(None, config, version="v2")
        else:
            inputs = {
                "task": message.content,
                "messages": [HumanMessage(content=message.content)],
                "revision_count": 0,
                "score": 0.0,
                "user_feedback": "",
                "search_queries": [],
                "dimension_scores": {},
                "questions_for_user": [],
            }
            stream = graph.astream_events(inputs, config, version="v2")

        # ── Session-scoped tracking ──────────────────────────────────
        active_steps: dict[str, cl.Step] = {}
        attempt = 1                          # writer/evaluator attempt counter

        async for event in stream:
            kind = event["event"]
            name = event["name"]
            data = event["data"]

            # ── Node Start ───────────────────────────────────────────
            if kind == "on_chain_start" and name in MAJOR_NODES:
                run_id = event.get("run_id")
                processed = cl.user_session.get("processed_ids") or set()
                cl.user_session.set("processed_ids", processed)
                if run_id in processed:
                    continue
                processed.add(run_id)

                # Skip the no-op ask_user node in the UI
                if name == "ask_user":
                    continue

                # Parent steps directly under the user message
                if name in ("writer", "evaluator"):
                    attempt = 1

                # Choose label
                if name == "writer":
                    label = _writer_label(attempt)
                elif name == "evaluator":
                    label = _evaluator_label(attempt)
                else:
                    label = _STEP_LABELS.get(name, name.capitalize())

                step = await _make_step(label, message.id)
                active_steps[name] = step

            # ── LLM Streaming ────────────────────────────────────────
            elif kind == "on_chat_model_stream":
                chunk_content = data["chunk"].content
                if not chunk_content:
                    continue
                node = event.get("metadata", {}).get("langgraph_node")
                if node and node in active_steps:
                    await active_steps[node].stream_token(chunk_content)

            # ── Node End ─────────────────────────────────────────────
            elif kind == "on_chain_end" and name in MAJOR_NODES:
                # Handle ask_user node — show planner questions
                if name == "ask_user":
                    final_state = graph.get_state(config)
                    questions = final_state.values.get("questions_for_user", [])
                    if questions:
                        await _show_planner_questions(questions)
                        # Mark that the next user message is an answer to planner
                        # questions, not researcher feedback.
                        cl.user_session.set("awaiting_planner_answers", True)
                    continue

                if name not in active_steps:
                    continue

                step = active_steps.pop(name)
                output = data.get("output")

                if name == "planner":
                    step.output = f"**Plan Generated**\n\n{output.get('plan', '')}"
                    await _end_step(step)

                elif name == "researcher":
                    await _finish_researcher(step, output)

                elif name == "writer":
                    step.output = output.get("draft", "")
                    await _end_step(step)

                elif name == "evaluator":
                    await _finish_evaluator(step, output)
                    attempt += 1

                elif name == "output":
                    await _finish_output(step, output)

        # ── After stream: display the final proposal ─────────────────
        final_state = graph.get_state(config)
        draft = final_state.values.get("draft", "")

        # Only show when graph has fully completed (no pending interrupt)
        if draft and not final_state.next:
            elements = [
                cl.File(
                    name="proposal_final.md",
                    content=draft.encode("utf-8"),
                    display="inline",
                ),
            ]

            # Generate PDF — don't let export errors block the message
            try:
                elements.append(cl.File(
                    name="proposal_final.pdf",
                    content=_draft_to_pdf_bytes(draft),
                    display="inline",
                ))
            except Exception as e:
                print(f"[WARNING] PDF generation failed: {e}")

            # Generate DOCX
            try:
                elements.append(cl.File(
                    name="proposal_final.docx",
                    content=_draft_to_docx_bytes(draft),
                    display="inline",
                ))
            except Exception as e:
                print(f"[WARNING] DOCX generation failed: {e}")

            await cl.Message(
                content=f"# 📄 Final Proposal\n\n{draft}\n\n---\n📥 **Download your proposal:** Use the attachments below to download as **PDF** or **DOCX**.",
                elements=elements,
            ).send()
    finally:
        cl.user_session.set("is_processing", False)


# ── Download helpers ─────────────────────────────────────────────────

def _draft_to_pdf_bytes(draft: str) -> bytes:
    """Convert a markdown draft string to PDF bytes using markdown-pdf (PyMuPDF)."""
    import tempfile, os
    from markdown_pdf import MarkdownPdf, Section

    # Professional CSS for clean tables with clear borders
    # NOTE: MuPDF does not support % values or nth-child selectors
    css = """
    body { font-family: sans-serif; font-size: 11pt; line-height: 1.6; color: #1a1a1a; }
    h1 { font-size: 20pt; color: #1a1a1a; margin-top: 18pt; margin-bottom: 8pt; }
    h2 { font-size: 16pt; color: #2C3E50; margin-top: 16pt; margin-bottom: 6pt; }
    h3 { font-size: 13pt; color: #34495E; margin-top: 12pt; margin-bottom: 4pt; }
    p  { margin: 0 0 6pt 0; }
    hr { border: none; border-top: 1px solid #bbb; margin: 12pt 0; }

    /* ── Tables with clear borders ───────────────── */
    table {
        border-collapse: collapse;
        margin: 10pt 0 14pt 0;
        font-size: 10pt;
    }
    th {
        background-color: #2C3E50;
        color: #ffffff;
        font-weight: bold;
        text-align: left;
        padding: 7pt 10pt;
        border: 1pt solid #1a252f;
    }
    td {
        padding: 6pt 10pt;
        border: 1pt solid #999999;
        vertical-align: top;
    }

    code { font-family: monospace; font-size: 9.5pt; background-color: #f0f0f0; padding: 1pt 3pt; }
    pre  { background-color: #f0f0f0; padding: 8pt; font-size: 9.5pt; }
    strong { font-weight: bold; }
    em { font-style: italic; }
    """

    pdf = MarkdownPdf(toc_level=0)
    pdf.add_section(Section(draft, toc=False), user_css=css)
    pdf.meta["title"] = "Proposal"

    # markdown-pdf requires saving to a file; read bytes back
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        pdf.save(tmp_path)
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _draft_to_docx_bytes(draft: str) -> bytes:
    """Convert a markdown draft string to DOCX bytes using python-docx."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = DocxDocument()

    # Tune default style
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def _add_run_with_inline(para, text: str):
        """Parse **bold** / *italic* inline and add runs."""
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))")
        for m in pattern.finditer(text):
            if m.group(2):  # **bold**
                run = para.add_run(m.group(2))
                run.bold = True
            elif m.group(3):  # *italic*
                run = para.add_run(m.group(3))
                run.italic = True
            elif m.group(4):  # `code`
                run = para.add_run(m.group(4))
                run.font.name = "Courier New"
            elif m.group(5):  # plain text
                para.add_run(m.group(5))

    for line in draft.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("---"):
            document.add_paragraph("─" * 60)
        elif re.match(r"^[-*]\s", stripped):
            para = document.add_paragraph(style="List Bullet")
            _add_run_with_inline(para, stripped[2:])
        elif stripped == "":
            document.add_paragraph("")
        else:
            para = document.add_paragraph()
            _add_run_with_inline(para, stripped)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ── Planner questions helper ─────────────────────────────────────────
async def _show_planner_questions(questions: list[str]) -> None:
    """Display the planner's questions to the user for HITL input."""
    q_list = "\n".join(f"{i}. {q}" for i, q in enumerate(questions, 1))
    await cl.Message(
        content=(
            "## \u2753 A few quick questions before I proceed\n\n"
            f"{q_list}\n\n"
            "Please answer the questions above and I'll incorporate "
            "your input into the proposal."
        ),
    ).send()


# ── Node-end helpers ─────────────────────────────────────────────────
async def _finish_researcher(step: cl.Step, output: dict) -> None:
    """Finalise the researcher step, then send the HITL checkpoint."""
    research_data = output.get("research_data", "")
    search_queries = output.get("search_queries", [])

    step.output = f"**Research Complete**\n\n{research_data}"
    await _end_step(step)

    query_list = "\n".join(f"  {i}. {q}" for i, q in enumerate(search_queries, 1))
    query_section = (
        f"**🔎 Queries used ({len(search_queries)}):**\n{query_list}"
        if search_queries
        else ""
    )

    preview = research_data[:600].rstrip()
    if len(research_data) > 600:
        preview += "\n\n*(…see full brief in the step above)*"

    actions = [
        cl.Action(name="proceed",           value="proceed",    label="✅ Proceed to Write", payload={"value": "proceed"}),
        cl.Action(name="edit_requirements",  value="edit",       label="✏️ Edit Requirements", payload={"value": "edit"}),
        cl.Action(name="reresearch",         value="reresearch", label="🔄 Re-research",      payload={"value": "reresearch"}),
    ]

    await cl.Message(
        content=(
            "## ✨ Research Phase Complete\n\n"
            f"{query_section}\n\n"
            "---\n\n"
            f"### Research Preview\n{preview}\n\n"
            "---\n\n"
            "Review the findings above. You can **proceed**, **edit requirements**, "
            "or **re-run research** with different focus."
        ),
        actions=actions,
    ).send()


async def _finish_evaluator(step: cl.Step, output: dict) -> None:
    """Render the scorecard inside the evaluator step."""
    score = output.get("score", 0.0)
    critique = output.get("critique", "")
    dimension_scores = output.get("dimension_scores", {})
    revision_count = output.get("revision_count", 1)

    scorecard = _build_scorecard(dimension_scores, score, revision_count)

    if critique and score < QUALITY_THRESHOLD:
        scorecard += f"\n\n### 📝 Critique\n{critique}"

    if score < QUALITY_THRESHOLD and revision_count < MAX_ITERATIONS:
        scorecard += (
            f"\n\n---\n⚡ **Score {score}/10** is below the {QUALITY_THRESHOLD} threshold. "
            f"Revising draft (attempt {revision_count + 1}/{MAX_ITERATIONS})…"
        )
    elif score >= QUALITY_THRESHOLD:
        scorecard += "\n\n---\n🎉 **Excellent!** The proposal meets the quality bar."

    step.output = scorecard
    await _end_step(step)


async def _finish_output(step: cl.Step, output: dict) -> None:
    """Finalise the output step (proposal display happens after the stream)."""
    step.output = "Saved to disk."
    await _end_step(step)


# ── Action callbacks ─────────────────────────────────────────────────
@cl.action_callback("proceed")
async def on_proceed(action: cl.Action):
    if cl.user_session.get("is_processing"):
        return
    await action.remove()
    # Send a visible message so it has a valid ID for nesting steps
    msg = cl.Message(content="✅ Proceeding to write the proposal...", author="User")
    await msg.send()
    await main(msg)


@cl.action_callback("edit_requirements")
async def on_edit(action: cl.Action):
    if cl.user_session.get("is_processing"):
        return
    await action.remove()
    cl.user_session.set("intent", "edit")
    await cl.Message(
        content="✏️ **Please type your specific requirements below.**\n\nI'll incorporate them into the proposal draft."
    ).send()


@cl.action_callback("reresearch")
async def on_reresearch(action: cl.Action):
    if cl.user_session.get("is_processing"):
        return
    await action.remove()
    cl.user_session.set("intent", "reresearch")
    await cl.Message(
        content="🔄 **Tell me what to focus the new research on.**\n\nFor example: *\"Focus more on competitor pricing\"* or *\"Research the European market instead\"*."
    ).send()
